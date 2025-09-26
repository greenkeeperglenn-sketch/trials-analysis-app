import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
import re
from io import BytesIO
import numpy as np
from scipy import stats
import statsmodels.api as sm
from statsmodels.formula.api import ols
from itertools import combinations

# For Word export
from docx import Document
from docx.shared import Inches

st.set_page_config(layout="wide")
st.title("Assessment Data Explorer")

# ======================
# Sidebar Global Settings
# ======================
st.sidebar.header("Global Settings")

alpha_options = {
    "Fungicide (0.05)": 0.05,
    "Biologicals in lab (0.10)": 0.10,
    "Biologicals in field (0.15)": 0.15,
}
alpha_label = st.sidebar.radio("Significance level:", list(alpha_options.keys()))
alpha_choice = alpha_options[alpha_label]

# ======================
# Helpers
# ======================
def parse_sheet_label_to_date(label: str):
    for dayfirst in (True, False):
        dt = pd.to_datetime(label, errors="coerce", dayfirst=dayfirst)
        if pd.notna(dt):
            return dt
    return None

def chronological_labels(labels):
    pairs = [(lab, parse_sheet_label_to_date(lab)) for lab in labels]
    pairs_sorted = sorted(
        pairs, key=lambda x: (pd.isna(x[1]), x[1] if pd.notna(x[1]) else pd.Timestamp.max)
    )
    return [p[0] for p in pairs_sorted]

def generate_cld_overlap(means, mse, df_error, alpha, rep_counts, a_is_lowest=True):
    trts = list(means.index)
    letters = {t: set() for t in trts}
    nsd = pd.DataFrame(False, index=trts, columns=trts)
    for t in trts:
        nsd.loc[t, t] = True
    t_crit = stats.t.ppf(1 - alpha / 2, df_error) if df_error > 0 else np.nan
    for a, b in combinations(trts, 2):
        n1, n2 = rep_counts.get(a, 1), rep_counts.get(b, 1)
        if n1 > 0 and n2 > 0 and pd.notna(mse) and pd.notna(t_crit):
            lsd_pair = t_crit * np.sqrt(mse * (1 / n1 + 1 / n2))
            diff = abs(means[a] - means[b])
            if diff <= lsd_pair:
                nsd.loc[a, b] = True
                nsd.loc[b, a] = True
    order = means.sort_values(ascending=a_is_lowest).index
    groups = []
    next_letter_code = ord("a")
    for t in order:
        joined_any = False
        for g in groups:
            if all(nsd.loc[t, m] for m in g["members"]):
                letters[t].add(g["letter"])
                g["members"].append(t)
                joined_any = True
        if not joined_any:
            new_letter = chr(next_letter_code)
            groups.append({"letter": new_letter, "members": [t]})
            letters[t].add(new_letter)
            next_letter_code += 1
        changed = True
        while changed:
            changed = False
            for g in groups:
                for cand in trts:
                    if g["letter"] not in letters[cand]:
                        if all(nsd.loc[cand, m] for m in g["members"]):
                            letters[cand].add(g["letter"])
                            g["members"].append(cand)
                            changed = True
    return {t: "".join(sorted(v)) for t, v in letters.items()}, nsd

def safe_key(base, assess):
    safe = re.sub(r"\W+", "_", str(assess))
    return f"{base}_{safe}"

def add_dataframe_to_word(doc: Document, df: pd.DataFrame):
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for row in df.itertuples(index=False):
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = "" if (pd.isna(val) if isinstance(val, float) else False) else str(val)

# ======================
# Upload & Parse
# ======================
uploaded_file = st.sidebar.file_uploader("Upload Excel File", type=["xlsx"])
if uploaded_file:
    xls = pd.ExcelFile(uploaded_file)
    all_data = []
    for sheet in xls.sheet_names:
        try:
            preview = pd.read_excel(xls, sheet_name=sheet, nrows=20)
            header_row = None
            for i, row in preview.iterrows():
                values = [str(v).lower() for v in row.values if pd.notna(v)]
                if any("block" in v for v in values) and any("treat" in v for v in values):
                    header_row = i
                    break
            if header_row is None:
                continue
            df = pd.read_excel(xls, sheet_name=sheet, skiprows=header_row)
            df.columns = df.iloc[0]
            df = df.drop(df.index[0])
            df = df.dropna(axis=1, how="all")
            df.columns = [str(c).strip() for c in df.columns]
            col_map = {c: re.sub(r"\W+", "", c).lower() for c in df.columns}
            block_col = next((o for o, n in col_map.items() if "block" in n), None)
            plot_col = next((o for o, n in col_map.items() if "plot" in n), None)
            treat_col = next((o for o, n in col_map.items() if "treat" in n or "trt" in n), None)
            if not (block_col and treat_col):
                continue
            treat_idx = df.columns.get_loc(treat_col)
            assess_list = df.columns[treat_idx + 1:].tolist()
            id_vars = [block_col, treat_col]
            if plot_col:
                id_vars.append(plot_col)
            df_long = df.melt(
                id_vars=id_vars,
                value_vars=assess_list,
                var_name="Assessment",
                value_name="Value",
            )
            df_long = df_long.rename(columns={block_col: "Block", treat_col: "Treatment"})
            df_long["DateLabel"] = sheet
            df_long["DateParsed"] = parse_sheet_label_to_date(sheet)
            all_data.append(df_long)
        except Exception:
            continue

    if not all_data:
        st.error("No valid tables found in this file.")
    else:
        data = pd.concat(all_data, ignore_index=True)

        # Treatments in sidebar
        treatments = sorted(data["Treatment"].dropna().unique(), key=lambda x: str(x))
        st.sidebar.subheader("Treatment Names")
        names_input = st.sidebar.text_area("Paste treatment names (one per line)", height=200)
        if names_input.strip():
            pasted = [n.strip() for n in names_input.split("\n") if n.strip()]
            if len(pasted) == len(treatments):
                mapping = dict(zip(treatments, pasted))
                data["Treatment"] = data["Treatment"].map(mapping).fillna(data["Treatment"])
                treatments = pasted
            else:
                st.sidebar.warning("Number of names pasted does not match detected treatments!")

        # Blocks in sidebar
        blocks = sorted(data["Block"].dropna().unique()) if "Block" in data.columns else []
        sel_blocks = st.sidebar.multiselect("Include Blocks", blocks, default=blocks, key="global_blocks")

        # Assessments in sidebar
        assessments_all = sorted(set(data["Assessment"].unique()))
        selected_assessments = st.sidebar.multiselect("Select assessments", assessments_all)

        # Date ordering
        date_labels_all = data["DateLabel"].dropna().unique().tolist()
        date_labels_ordered = chronological_labels(date_labels_all)

        # Apply global block filter
        if sel_blocks:
            data = data[data["Block"].isin(sel_blocks)]

        color_map = {
            t: px.colors.qualitative.Plotly[i % len(px.colors.qualitative.Plotly)]
            for i, t in enumerate(treatments)
        }

        all_tables = {}
        word_doc = Document()
        word_doc.add_heading("Assessment Report", level=1)

        # ======================
        # Loop through assessments
        # ======================
        for assess in selected_assessments:
            with st.expander(f"Assessment: {assess}", expanded=False):
                st.markdown(f"<h2 style='text-align:center'>{assess}</h2>", unsafe_allow_html=True)

                # ----------------------
                # Chart Settings
                # ----------------------
                with st.expander("Chart Settings", expanded=False):
                    view_mode_chart = st.radio(
                        "Grouping",
                        ["By Date", "By Treatment"],
                        key=safe_key("viewmode", assess),
                    )
                    a_is_lowest_chart = (
                        st.radio(
                            "Lettering convention",
                            ["Lowest = A", "Highest = A"],
                            index=0,
                            key=safe_key("letters_mode", assess),
                        )
                        == "Lowest = A"
                    )
                    axis_min = st.number_input(
                        "Y-axis minimum",
                        value=int(data["Value"].min()) if not data["Value"].empty else 0,
                        step=1,
                        key=safe_key("ymin", assess),
                    )
                    axis_max = st.number_input(
                        "Y-axis maximum",
                        value=int(data["Value"].max()) if not data["Value"].empty else 100,
                        step=1,
                        key=safe_key("ymax", assess),
                    )
                    chart_mode = st.radio(
                        "Chart type",
                        ["Boxplot", "Bar chart"],
                        key=safe_key("chartmode", assess),
                    )
                    add_se = add_lsd = add_letters = False
                    if chart_mode == "Bar chart":
                        add_se = st.checkbox("Add SE error bars", key=safe_key("se", assess))
                        add_lsd = st.checkbox("Add LSD error bars", key=safe_key("lsd", assess))
                        add_letters = st.checkbox("Add statistical letters", key=safe_key("letters", assess))

                # Prepare filtered data for this assessment
                df_sub = data[data["Assessment"] == assess].copy()
                df_sub["Value"] = pd.to_numeric(df_sub["Value"], errors="coerce")
                df_sub = df_sub.dropna(subset=["Value"])
                df_sub["DateLabel"] = pd.Categorical(
                    df_sub["DateLabel"], categories=date_labels_ordered, ordered=True
                )
                visible_treatments = st.multiselect(
                "Show treatments (does not recalc stats)",
                options=treatments,
                default=treatments,
                key=safe_key("visible_treatments", assess),
                )
                # ----------------------
                # Chart
                # ----------------------
                with st.expander("Chart", expanded=True):
                    if chart_mode == "Boxplot":
                        df_plot = df_sub[df_sub["Treatment"].isin(visible_treatments)].copy()
                        if view_mode_chart == "By Date":
                            fig = px.box(
                                df_plot,
                                x="DateLabel",
                                y="Value",
                                color="Treatment",
                                color_discrete_map=color_map,
                                category_orders={"DateLabel": date_labels_ordered, "Treatment": treatments},
                            )
                        else:
                            fig = px.box(
                                df_plot,
                                x="Treatment",
                                y="Value",
                                color="DateLabel",
                                category_orders={"Treatment": treatments, "DateLabel": date_labels_ordered},
                            )
                        fig.update_traces(boxpoints=False)
                    else:
                        agg = (
                            df_sub.groupby(["DateLabel", "Treatment"])["Value"]
                            .agg(mean="mean", count="count", std="std")
                            .reset_index()
                        )
                        agg["se"] = agg["std"] / np.sqrt(agg["count"])
                        letters_dict, lsd_by_date = {}, {}
                        for date_label in list(df_sub["DateLabel"].cat.categories):
                            df_date = df_sub[df_sub["DateLabel"] == date_label]
                            if df_date["Treatment"].nunique() > 1 and len(df_date) > 1:
                                try:
                                    if "Block" in df_date.columns:
                                        model = ols("Value ~ C(Treatment) + C(Block)", data=df_date).fit()
                                    else:
                                        model = ols("Value ~ C(Treatment)", data=df_date).fit()
                                    anova = sm.stats.anova_lm(model, typ=2)
                                    df_error = float(model.df_resid)
                                    mse = float(anova.loc["Residual", "sum_sq"] / df_error)
                                    means_date = df_date.groupby("Treatment")["Value"].mean()
                                    rep_counts_date = df_date["Treatment"].value_counts().to_dict()
                                    letters, _ = generate_cld_overlap(
                                        means_date, mse, df_error, alpha_choice, rep_counts_date,
                                        a_is_lowest=a_is_lowest_chart
                                    )
                                    letters_dict[date_label] = letters
                                    n_avg = np.mean(list(rep_counts_date.values()))
                                    lsd_val = stats.t.ppf(1 - alpha_choice/2, df_error) * np.sqrt(2*mse/n_avg)
                                    lsd_by_date[date_label] = lsd_val
                                except Exception:
                                    letters_dict[date_label] = {}
                                    lsd_by_date[date_label] = np.nan
                            else:
                                letters_dict[date_label] = {}
                                lsd_by_date[date_label] = np.nan
                        agg["letters"] = agg.apply(
                            lambda r: letters_dict.get(r["DateLabel"], {}).get(r["Treatment"], ""), axis=1
                        )
                        agg["LSD"] = agg["DateLabel"].map(lsd_by_date)
                        if view_mode_chart == "By Date":
                            x_axis, group_axis = "DateLabel", "Treatment"
                            category_orders_bar = {"DateLabel": list(df_sub["DateLabel"].cat.categories), "Treatment": treatments}
                        else:
                            x_axis, group_axis = "Treatment", "DateLabel"
                            category_orders_bar = {"Treatment": treatments, "DateLabel": list(df_sub["DateLabel"].cat.categories)}
                        value_col = "mean"
                        fig = go.Figure()
                        x_order = category_orders_bar[x_axis]
                        agg_plot = agg[agg["Treatment"].isin(visible_treatments)].copy()
                        for group in agg_plot[group_axis].dropna().unique():
                            df_g = agg_plot[agg_plot[group_axis] == group].copy()
                            df_g = df_g.set_index(x_axis).reindex(x_order).reset_index()
                            error_array = None
                            if add_se:
                                error_array = df_g["se"].to_numpy()
                            elif add_lsd:
                                error_array = df_g["LSD"].to_numpy()
                            texts = df_g["letters"].tolist() if add_letters else None
                            fig.add_trace(go.Bar(
                                x=df_g[x_axis],
                                y=df_g[value_col],
                                name=str(group),
                                error_y=dict(type="data", array=error_array, visible=True) if error_array is not None else None,
                                text=texts,
                                textposition="outside" if add_letters else None,
                                textfont=dict(color="black", size=12) if add_letters else None,
                            ))
                        fig.update_layout(barmode="group", xaxis=dict(categoryorder="array", categoryarray=x_order))
                    fig.update_yaxes(range=[axis_min, axis_max])
                    st.plotly_chart(fig, use_container_width=True)

                # ----------------------
                # Treatment Filter
                # ----------------------
                visible_treatments = st.multiselect(
                    "Show treatments (does not recalc stats)",
                    options=treatments,
                    default=treatments,
                    key=safe_key("visible_treatments", assess),
                )
                # ----------------------
                # Statistics Table
                # ----------------------
                with st.expander("Statistics Table", expanded=True):
                    wide_table = pd.DataFrame({"Treatment": treatments})
                    summaries = {}
                
                    for date_label in date_labels_ordered:
                        df_date = df_sub[df_sub["DateLabel"] == date_label].copy()
                        if df_date.empty:
                            wide_table[f"{date_label}"] = np.nan
                            wide_table[f"{date_label} S"] = ""
                            summaries[date_label] = {"P": np.nan, "LSD": np.nan, "d.f.": np.nan, "%CV": np.nan}
                            continue
                
                        df_date["Value"] = pd.to_numeric(df_date["Value"], errors="coerce")
                        df_date = df_date.dropna(subset=["Value"])
                
                        if df_date["Treatment"].nunique() > 1 and len(df_date) > 1:
                            means = df_date.groupby("Treatment")["Value"].mean()
                            rep_counts = df_date["Treatment"].value_counts().to_dict()
                
                            try:
                                if "Block" in df_date.columns:
                                    model = ols("Value ~ C(Treatment) + C(Block)", data=df_date).fit()
                                else:
                                    model = ols("Value ~ C(Treatment)", data=df_date).fit()
                                anova = sm.stats.anova_lm(model, typ=2)
                                df_error = float(model.df_resid)
                                mse = float(anova.loc["Residual", "sum_sq"] / df_error)
                                p_val = float(anova.loc["C(Treatment)", "PR(>F)"])
                            except Exception:
                                df_error, mse, p_val = np.nan, np.nan, np.nan
                
                            cv = 100 * np.sqrt(mse) / means.mean() if pd.notna(mse) and means.mean() != 0 else np.nan
                
                            letters, _ = generate_cld_overlap(
                                means, mse, df_error, alpha_choice, rep_counts, a_is_lowest=a_is_lowest_chart
                            )
                
                            n_avg = np.mean(list(rep_counts.values()))
                            lsd_val = (
                                stats.t.ppf(1 - alpha_choice / 2, df_error) * np.sqrt(2 * mse / n_avg)
                                if pd.notna(mse) else np.nan
                            )
                
                            wide_table[f"{date_label}"] = wide_table["Treatment"].map(means)
                            wide_table[f"{date_label} S"] = wide_table["Treatment"].map(letters).fillna("")
                            summaries[date_label] = {"P": p_val, "LSD": lsd_val, "d.f.": df_error, "%CV": cv}
                
                        else:
                            wide_table[f"{date_label}"] = np.nan
                            wide_table[f"{date_label} S"] = ""
                            summaries[date_label] = {"P": np.nan, "LSD": np.nan, "d.f.": np.nan, "%CV": np.nan}
                
                    # Add summary rows
                    summary_rows = []
                    for metric in ["P", "LSD", "d.f.", "%CV"]:
                        row = {"Treatment": metric}
                        for date_label in date_labels_ordered:
                            row[f"{date_label}"] = summaries[date_label][metric]
                            row[f"{date_label} S"] = ""
                        summary_rows.append(row)
                
                    wide_table = pd.concat([wide_table, pd.DataFrame(summary_rows)], ignore_index=True)
                    wide_table = wide_table.round(2)
                
                    st.dataframe(wide_table, use_container_width=True, hide_index=True)
                
                    all_tables[assess] = wide_table

               

                    # Append summary rows (always based on full data)
                    summary_rows = []
                    for metric in ["P", "LSD", "d.f.", "%CV"]:
                        row = {"Treatment": metric}
                        for date_label in date_labels_ordered:
                            row[f"{date_label}"] = summaries[date_label][metric]
                            row[f"{date_label} S"] = ""
                        summary_rows.append(row)

                    wide_table = pd.concat([wide_table, pd.DataFrame(summary_rows)], ignore_index=True)
                    wide_table = wide_table.round(1)

                    # Filtered view for display: show only selected treatments + summary rows
                    summary_labels = {"P", "LSD", "d.f.", "%CV"}
                    mask_visible = wide_table["Treatment"].isin(visible_treatments) | wide_table["Treatment"].isin(summary_labels)
                    wide_table_view = wide_table[mask_visible].reset_index(drop=True)

                    st.dataframe(
                        wide_table_view,
                        use_container_width=True,
                        hide_index=True,
                        column_config={"Treatment": st.column_config.Column(pinned="left")},
                    )

                    # Keep the full table for exports
                    all_tables[assess] = wide_table

                # ----------------------
                # Append this assessment to Word (chart as currently shown + full table)
                # ----------------------
                word_doc.add_heading(f"Assessment: {assess}", level=2)
                try:
                    img_buffer = BytesIO()
                    pio.write_image(fig, img_buffer, format="png", scale=2)  # requires 'kaleido'
                    img_buffer.seek(0)
                    word_doc.add_picture(img_buffer, width=Inches(6.5))
                    word_doc.add_paragraph()
                except Exception:
                    st.warning(
                        f"Could not export chart image for '{assess}' to Word (need 'kaleido'). "
                        f"Tables will still be exported."
                    )
                # Add the full (unfiltered) table to Word
                table = word_doc.add_table(rows=1, cols=len(wide_table.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(wide_table.columns):
                    hdr_cells[i].text = str(col)
                for row in wide_table.itertuples(index=False):
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = "" if (isinstance(val, float) and pd.isna(val)) else str(val)
                word_doc.add_paragraph()

        # ----------------------
        # Download buttons (Excel: full tables; Word: charts as shown + full tables)
        # ----------------------
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
            for assess, table in all_tables.items():
                table.round(1).to_excel(writer, sheet_name=assess[:30], index=False)
        st.download_button(
            "Download Tables (Excel)",
            data=buffer,
            file_name="assessment_tables.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        word_buffer = BytesIO()
        word_doc.save(word_buffer)
        word_buffer.seek(0)
        st.download_button(
            "Download Report (Word)",
            data=word_buffer,
            file_name="assessment_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )



